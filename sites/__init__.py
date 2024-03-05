#!/usr/bin/env python
# -*- coding: utf-8 -*-
import os
from flask import Flask, request, render_template, make_response, redirect, session,flash,send_file
from flask_sqlalchemy import SQLAlchemy
import datetime
from flask import request, make_response, json
import hashlib
from cryptography.fernet import Fernet
import pandas as pd
import io
import docx
import re
import cv2
from flask import request
import base64
import requests
import numpy as np
from PIL import Image
import pymysql
# 静态文件路径配置
app = Flask(__name__, template_folder='../templates', static_folder='../templates/layui')
# 读取config配置文件
conf_path = os.path.join(
    os.path.abspath(os.path.dirname(__file__)), "../config.py")
app.config.from_pyfile(conf_path)
# 初始化mysql连接
db = SQLAlchemy(app)


@app.route("/")
def index():
    """
    主路由
    :return:
    """
    # 判断是否有登录状态 如果有的话跳转home否则到login
    try:
        session['user']
    except Exception:
        return render_template('login.html')

    user_type = session['user_type']
    return render_template('home.html', user_type=user_type, user=session['user'])


@app.route("/home")
def home():
    # 判断是否有登录状态 如果有的话跳转home否则到login
    try:
        session['user']
    except Exception:
        return redirect('/')
    user_type = session['user_type']
    return render_template('home.html', user_type=user_type, user=session['user'])


def register_blueprint():
    from sites.user.views import user
    from sites.admin.views import admin
    from sites.file.views import file
    from sites.open_file.views import open_file
    from sites.teacher.views import teacher
    from sites.student.views import student
    app.register_blueprint(user)
    app.register_blueprint(student)
    app.register_blueprint(teacher)
    app.register_blueprint(admin)
    app.register_blueprint(file)
    app.register_blueprint(open_file)
@app.route("/self")
def self():
    return render_template('pw.html')
@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        flash('No file uploaded')
        return redirect('/')
    file = request.files['file']
    if file.filename == '':
        flash('No file selected')
        return redirect('/')
    if file:
        text1 = request.form['text1']
        text2 = request.form['text2']
        text3 = request.form['text3']
        df = pd.read_excel(file)
        btn_value = request.form['btn']
        if btn_value == '1':
            # 定义密钥，可以使用Fernet库生成
            key = b'qm2B_jw-IZ8SKR5zAE88i0-T0X4x4HPRt--5b5i5qec='

            # 创建Fernet对象
            f = Fernet(key)
            encrypted_text1 = f.encrypt(text1.encode())
            encrypted_text2 = f.encrypt(text2.encode())
            encrypted_text3 = f.encrypt(text3.encode())
            # 将加密后的密文替换原始数据
            df[text1] = encrypted_text1.decode()
            df[text2] = encrypted_text2.decode()
            df[text3] = encrypted_text3.decode()
        elif btn_value == '2':
            hash1 = hashlib.sha256(text1.encode()).hexdigest()
            hash2 = hashlib.sha256(text2.encode()).hexdigest()
            hash3 = hashlib.sha256(text3.encode()).hexdigest()
            # 将加密后的哈希值替换原始数据
            df[text1] = hash1
            df[text2] = hash2
            df[text3] = hash3
        else:
            # 如果没有选择任何按钮，则返回错误页面
            return render_template('database.html', message='请选择一个按钮')
        # 在这里添加数据脱敏的代码，比如将姓名和地址替换为随机字符串
        #df[text1] = '*****'
        #df[text2] = '*****'
        #df[text3] = '12345'
        # 将修改后的DataFrame转换为bytesIO对象，以便将其发送给用户
        output = io.BytesIO()
        writer = pd.ExcelWriter(output, engine='xlsxwriter')
        df.to_excel(writer, index=False, sheet_name='Sheet1')
        writer.close()
        output.seek(0)
        return send_file(output, download_name='output.xlsx', as_attachment=True)

@app.route('/word_change', methods=['POST'])
def word_change():
        # 定义电话号码匹配的正则表达式
        file = request.files['file']
        # Read the contents of the file
        file_contents = file.read()
        # Create a file-like object from the bytes object
        file_obj = io.BytesIO(file_contents)

        # Convert the file contents to a docx.Document object
        doc = docx.Document(file_obj)
        search_text = request.form['search_text']
        # Convert the file contents to a docx.Document object
        # doc = docx.Document(file_contents)
        # doc = docx.Document('新建 DOCX 文档.docx')
        text = '\n'.join([p.text for p in doc.paragraphs])  # 将文档中所有段落的文本内容合并为一个字符串
        iphone = r'(电话|联系方式|电话号码|联系电话)'
        phone_pattern = re.compile(r'\(?0\d{2,3}[)-]?\d{7,8}|1[3456789]\d{9}')
        name_pattern = re.compile(r'[\u4e00-\u9fa5]{2,4}')
        id_pattern = re.compile(r'\d{17}[\dXx]')
        address_pattern = re.compile(
            r'[\u4e00-\u9fa5]{3,15}(省|市|自治区|自治州|县|区|镇|乡|街道)[\u4e00-\u9fa5]{2,20}(号|弄|巷|路|街|村)')
        score_pattern = re.compile(r'\d{1,3}\.\d{1,2}|\d{1,3}')
        names = []
        phones = []
        ids = []
        addresses = []
        scores = []
        for match in name_pattern.finditer(text):
            name = match.group()
            start = match.start()
            end = match.end()

            if len(name) == 2 and (name[0] in ['王', '李', '张', '刘']):
                names.append(name)
            elif len(name) == 3 and (name[0] in ['赵', '钱', '孙', '李', '周', '吴', '郑', '王']):
                names.append(name)
            else:
                context = text[max(0, start - 2):end + 2]
                if ('先生' in context) or ('女士' in context) or ('同学' in context):
                    names.append(name)

        for phone in phone_pattern.findall(text):
            phones.append(phone)
        for id_number in id_pattern.findall(text):
            ids.append(id_number)

        for address in address_pattern.findall(text):
            addresses.append(address)

        for score in score_pattern.findall(text):
            scores.append(score)
        # 将姓名和电话号码替换为指定的内容
        if re.match(iphone, search_text):
            # 如果搜索文本是电话号码，则将所有匹配的电话号码替换为指定的内容
            for phone in phones:
                text = text.replace(phone, '000-0000-0000')
        elif search_text == "姓名":
            # 如果搜索文本是姓名，则将所有匹配的姓名替换为指定的内容
            for name in names:
                text = text.replace(name, 'XXX')
        elif search_text == "家庭住址":
            for address in addresses:
                text = text.replace(address, 'xx省xx市xx')
        elif search_text == "考试成绩":
            for score in scores:
                text = text.replace(score, 'xx')

        # 生成新的文档
        doc = docx.Document()
        doc.add_paragraph(text)
        file_dir = os.path.join(app.root_path, 'templates', 'layui', 'temp')
        if not os.path.exists(file_dir):
            os.makedirs(file_dir)
        file_path = os.path.join(file_dir, 'new.docx')

        # 保存 Word 文档到指定路径
        doc.save(file_path)

        # 下载文件
        return send_file(file_path, as_attachment=True)

@app.route('/picture_change',methods=['POST'])
def picture_change():
    file = request.files['file']
    # 使用Pillow库打开图片文件
    #file_str =str(file)
    #img = cv2.imread('static/images/bj.jpg')
    # 加载图片
    #file = request.files['file']
    #f = open(file, 'rb')
    # 将文件读取到内存中
    file_data = np.frombuffer(file.read(), np.uint8)
    #f = open('static/images/lan.jpg', 'rb')
    #img = base64.b64encode(f.read())

    #params = {"image": img}
    # 使用OpenCV读取图像数据
    img = cv2.imdecode(file_data, cv2.IMREAD_COLOR)
    #img = cv2.imencode('.jpg', img)[1].tostring()
    # 从POST请求中获取上传的文件
    api_url = "https://aip.baidubce.com/rest/2.0/ocr/v1/general"
    app_id = "	32232315"
    api_key = "AyZ03QI6fn0Kkx6AwlH6Zv3p"
    secret_key = "Di3HMePOVUjRxw98c6NhMmrres9f0eiB"
    host = 'https://aip.baidubce.com/oauth/2.0/token?grant_type=client_credentials&client_id=AyZ03QI6fn0Kkx6AwlH6Zv3p&client_secret=Di3HMePOVUjRxw98c6NhMmrres9f0eiB'
    response = requests.get(host)
    if response:
        # print(response.json()) # 这是返回的字段
        # print('token：', response.json()['refresh_token'])
        # print('有效期：', response.json()['expires_in'])
        access_token = response.json()['access_token']
    #client = AipImageClassify(app_id, api_key, secret_key)
    api_url = api_url + "?access_token=" + access_token
    headers = {"Content-Type": "application/x-www-form-urlencoded"}
    params = {"access_token": access_token, "language_type": "CHN_ENG", "detect_direction": "true",
              "recognize_granularity": "small"}
    # 加载图片
    # 从POST请求中获取上传的文件
    #params = {"image": img}

    # 读取图像并调用百度API进行OCR识别
    img_str = cv2.imencode('.jpg', img)[1].tostring()
    params['image'] = str(base64.b64encode(img_str), encoding='utf-8')
    try:
        res = requests.post(api_url, headers=headers, data=params, timeout=10)  # 超时设置为10秒
    except:
        for i in range(4):  # 循环去请求网站
            res = requests.post(api_url, headers=headers, data=params,  timeout=20)
            if res.status_code == 200:
                break
    #res = requests.post(API_url, headers=headers, params=params, timeout=500)
    #res = res.read().decode()
    #print(res.text)
    print(res.json())
    result = json.loads(res.text)
    #json_str = json.dumps(result['words_result'])
    #print(json.dumps(result, indent=4, ensure_ascii=False))
    # Convert JSON string to dictionary
    #dict_obj = json.loads(json_str)
    search_text = request.form['search_text']
    # 显示结果图像
    for i, word in enumerate(result['words_result']):
        text = word['words']
        if search_text in text:
            loc = word['location']
            x, y = loc['left'], loc['top']
            w, h = loc['width'], loc['height']
            # 在图像上画矩形框
            roi = img[y:y + h, x:x + w]
            blurred = cv2.GaussianBlur(roi, (51, 51), 0)
            img[y:y + h, x:x + w] = blurred
    cv2.imshow("result", img)
    cv2.imwrite('output.jpg', img)
    cv2.waitKey(0)
    cv2.destroyAllWindows()

    parent_dir = os.path.abspath(os.path.join(os.getcwd(), os.pardir))

    # 定义保存文件的路径（位于上一层目录的temp文件夹）
    file_dir = os.path.join(parent_dir, 'file_system')
    if not os.path.exists(file_dir):
        os.makedirs(file_dir)

    # 假设要保存的文件名为 'new.jpg'
    original_file_path = os.path.join(file_dir, 'output.jpg')

    # 打开原始图片
    image = Image.open(original_file_path)

    # 保存图片到不同的文件名（例如 'output.jpg'）以避免覆盖原始文件
    output_file_path = os.path.join(file_dir, 'output.jpg')
    image.save(output_file_path)

    # 将保存的图片作为文件附件发送
    return send_file(output_file_path, as_attachment=True)
register_blueprint()
